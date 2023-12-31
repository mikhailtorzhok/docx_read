import docx
import sys
import os
import config
from selenium import webdriver
from selenium.webdriver.firefox.options import Options
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import Select
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys
import time 
import signal
import stopit
import pandas as pd

timedelay=0.2
#Ur_list = []
#Fiz_list = []
#Trash_list = []
#Ur_dictionary = {}
#Fiz_dictionary = {}
#Trash_dictionary = {}
time606=3

def main():
    print("Hello World!")   
    delay = 5 # seconds
    driver = chrome_settings_func()
    
    #site_login_func(driver, delay)

    while True:
   

        read_from_word_and_write_to_site(driver, delay, 'Doc1.docx')
        break



def  read_from_word_and_write_to_site(driver, delay, filename='Doc1.docx'):
    
    doc = docx.Document(filename)
    
    
    
    
    #print not empty tables text
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if bool(cell.text):
                    #print(cell.text)
                    table_text.append(cell.text)
                    
    #print body text
    header_text = []
    body_text = []
    i=0
    previous_text_was_bold = True
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if not run.bold:
                #print(i)
                #print(run.text)
                if previous_text_was_bold:
                    body_text.insert(i,run.text)
                else:
                    body_text[i] = body_text[i] + os.linesep + run.text
                previous_text_was_bold = False
            else:
                #print(i)
                # print(run.text)
                if not previous_text_was_bold:
                    #print('i=' + str(i))
                    #rint(header_text)
                    i+=1
                    header_text.insert(i,run.text)              
                    
                else:
                    if i==0:
                        header_text.insert(i,run.text)
                        #header_text.insert(i,header_text[i]+run.text)
                    else:
                        header_text[i] = header_text[i] + os.linesep + run.text
                        #header_text.insert(i,header_text[i]+run.text)
                                  
                previous_text_was_bold = True
    
    '''
    #example to output
    print("----------------------###############----------------------------")
    print(header_text[6])
    print("----------------------###############----------------------------")
    print(table_text[6])
    print("----------------------###############----------------------------")
    print(body_text[6])
    print("----------------------###############----------------------------")
    '''
    
    result_text = []
    
    for j in range(i+1):
        if "\r\n" not in header_text[j]:
            continue
        else:
            print("WARNING carriage return was detected in text")
            #print(header_text[j])
            header_text[j] = ' '.join(header_text[j].splitlines())
            #print(header_text[j])
           
    '''
    for j in range(i+1):
        print("----------------------###############----------------------------")
        print(header_text[j])
        print("----------------------###############----------------------------")
        print(table_text[j])
        print("----------------------###############----------------------------")
        print(body_text[j])
        print("----------------------###############----------------------------")
    '''
  
         
    return



def chrome_settings_func():
    chrome_options = Options()
    
    chrome_options.add_experimental_option("detach", True)
    driver = webdriver.Chrome(options=chrome_options) 
    #driver.set_window_size(1920, 1080)
    #driver.get("https://torsed.voskhod.ru/app/#!")
    return driver
   
def site_login_func(driver, delay):
    login_input = WebDriverWait(driver, delay).until(EC.presence_of_element_located((By.XPATH, '//input')))
    print ("Page is ready!")
    login_input = driver.find_elements(By.XPATH, '//input')
    print(login_input) 
    login_input[0].send_keys(config.username)
    login_input[1].send_keys(config.password)
            
    button_login = driver.find_element(By.XPATH, '//div[@class="v-button v-widget cuba-login-submit v-button-cuba-login-submit v-has-width"]')
    button_login.click()

    return




def find_and_click_element_by_path(driver, delay, path):
    button = WebDriverWait(driver, delay).until(EC.presence_of_element_located((By.XPATH, path)))
    action = ActionChains(driver)
    action.move_to_element(button)
    action.click(button)
    action.perform()
    return







if __name__ == "__main__":
    main()

